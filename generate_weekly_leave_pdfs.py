import csv
import re
from pathlib import Path
from datetime import datetime, date
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE_DOCX = "상담팀_휴가신청서.docx"
CSV_PATH = "weekly.csv"

OUT_DIR = Path("output")
TMP_DIR = OUT_DIR / "_tmp_docx"

def sanitize_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r'[\\/:*?"<>|]', "_", s)   # 윈도우 금지문자 제거
    s = re.sub(r"\s+", " ", s)
    return s.strip() or "unknown"

def parse_ymd(s: str) -> date:
    s = (s or "").strip().replace("/", "-")
    return datetime.strptime(s, "%Y-%m-%d").date()

def label_leave_type(leave_type: str) -> str:
    if leave_type == "annual":
        return "휴가"
    if leave_type == "half_am":
        return "오전 반차"
    if leave_type == "half_pm":
        return "오후 반차"
    return "휴가"

def build_d_date(leave_type: str, start: date, end: date) -> str:
    # 요구사항:
    # - 하루만 신청: 2026-02-04
    # - 오전 반차: 2026-02-04 오전 반차
    # - 오후 반차: 2026-02-04 오후 반차
    if leave_type in ("half_am", "half_pm"):
        return f"{start.strftime('%Y-%m-%d')} {label_leave_type(leave_type)}"

    # 연차(휴가)
    if start == end:
        return start.strftime("%Y-%m-%d")
    return f"{start.strftime('%Y-%m-%d')}~{end.strftime('%Y-%m-%d')}"

def format_filename_range(start: date, end: date) -> str:
    # 파일명 괄호 표기:
    # - 단일: (26.01.30)
    # - 같은 달 연속: (26.01.29~30)
    # - 달 넘어감: (26.01.31~26.02.01)
    yy = start.strftime("%y")
    mm = start.strftime("%m")
    dd1 = start.strftime("%d")

    if start == end:
        return f"({yy}.{mm}.{dd1})"

    if start.year == end.year and start.month == end.month:
        dd2 = end.strftime("%d")
        return f"({yy}.{mm}.{dd1}~{dd2})"

    yy2 = end.strftime("%y")
    mm2 = end.strftime("%m")
    dd2 = end.strftime("%d")
    return f"({yy}.{mm}.{dd1}~{yy2}.{mm2}.{dd2})"

def extract_number_from_email(email: str) -> str:
    email = (email or "").strip()
    if "@" not in email:
        return ""
    return email.split("@", 1)[0].strip()

def replace_tokens_in_paragraph(paragraph, mapping: dict):
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return

    new_text = full
    for k, v in mapping.items():
        new_text = new_text.replace(f"{{{{{k}}}}}", str(v))

    if new_text != full:
        paragraph.clear()
        run = paragraph.add_run(new_text)

        # ===== 폰트 강제 지정 =====
        run.font.name = "Malgun Gothic"   # 맑은 고딕
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(11)

def replace_tokens_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_tokens_in_paragraph(p, mapping)
            for t in cell.tables:
                replace_tokens_in_table(t, mapping)

def fill_docx(template_path: str, mapping: dict, out_docx_path: Path):
    doc = Document(template_path)
    for p in doc.paragraphs:
        replace_tokens_in_paragraph(p, mapping)
    for t in doc.tables:
        replace_tokens_in_table(t, mapping)
    doc.save(out_docx_path)

def convert_docx_to_pdf_docx2pdf(docx_path: Path, pdf_path: Path):
    from docx2pdf import convert
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    convert(str(docx_path), str(pdf_path))

def main():
    OUT_DIR.mkdir(exist_ok=True)
    TMP_DIR.mkdir(exist_ok=True)

    with open(CSV_PATH, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    if not rows:
        print("CSV에 데이터가 없습니다.")
        return

    for r in rows:
        name = (r.get("name") or "").strip()
        email = (r.get("email") or "").strip()
        number = extract_number_from_email(email)

        leave_type = (r.get("leave_type") or "annual").strip()
        start = parse_ymd(r.get("start_date"))
        end = parse_ymd(r.get("end_date") or r.get("start_date"))

        submitted_raw = (r.get("submitted_at") or "").strip()
        submitted = parse_ymd(submitted_raw) if submitted_raw else date.today()

        d_date = build_d_date(leave_type, start, end)

        mapping = {
            "name": name,
            "number": number,
            "d_date": d_date,
            "year": str(submitted.year),
            "month": str(submitted.month),
            "date": str(submitted.day),
        }

        range_part = format_filename_range(start, end)
        type_part = label_leave_type(leave_type)
        base_name = f"{range_part}{name} 팀장 {type_part}"
        safe_base = sanitize_filename(base_name)

        out_docx = TMP_DIR / f"{safe_base}.docx"
        out_pdf = OUT_DIR / f"{safe_base}.pdf"

        fill_docx(TEMPLATE_DOCX, mapping, out_docx)
        convert_docx_to_pdf_docx2pdf(out_docx, out_pdf)

        print("OK:", out_pdf)

    print("\n완료: output 폴더 확인")

if __name__ == "__main__":
    main()
